"""
build_db.py  —  Run once to parse all rate sheets into quotation.db

Usage:
    python features/quotation/build_db.py

Reads:
    data/BKK_PTY_Land_Rates_VA_Format__01_Nov_2025_to_31_Dec_2026.xlsx
    data/Good_Day_Vacation_-__Phuket_-_Krabi_-_Samui_-_Phang_Nga_-_31_Oct_2026.xlsx
    data/PHUKET_CONTRACT__RATES_2026___VAYOAURA___UPDATED_11_JAN_26.docx
    data/KRABI_CONTRACT_RATES__2026.docx
    data/SAMUI_CONTRACT_RATES_2026___NO_VAT___UPDATED_24_MAR_26.docx

Outputs:
    features/quotation/quotation.db   (SQLite)
"""

import re
import sqlite3
from pathlib import Path

import pandas as pd
from docx import Document

# ── paths ────────────────────────────────────────────────────────────────────
BASE       = Path(__file__).parent
DATA       = BASE.parent.parent / "data"       # adjust if you keep files elsewhere
DB_PATH    = BASE / "quotation.db"

BKK_XL    = DATA / "BKK_PTT.xlsx"
GDV_XL    = DATA / "GOOD_DAY.xlsx"
PHK_DOCX  = DATA / "PHUKET_CR.docx"
KRB_DOCX  = DATA / "KRABI_CR.docx"
SMU_DOCX  = DATA / "SAMUI_CR.docx"


# ── helpers ──────────────────────────────────────────────────────────────────
def to_int(val) -> int | None:
    """Safely coerce a value to int, return None if not possible."""
    try:
        v = str(val).strip().replace(",", "").split("/")[0]
        return int(float(v))
    except Exception:
        return None


def clean(val: str) -> str:
    """Strip whitespace and normalise newlines."""
    return re.sub(r"\s+", " ", str(val or "")).strip()


def unique_cell(row_cells: list[str]) -> str:
    """
    Word tables have merged cells that repeat the same text across columns.
    Return the first unique non-empty value in the row of deduplicated cells.
    """
    seen = []
    for c in row_cells:
        t = clean(c)
        if t and t not in seen:
            seen.append(t)
    return seen[0] if seen else ""


# ── DB setup ─────────────────────────────────────────────────────────────────
def init_db(conn: sqlite3.Connection):
    conn.executescript("""
    CREATE TABLE IF NOT EXISTS services (
        id           INTEGER PRIMARY KEY AUTOINCREMENT,
        source       TEXT NOT NULL,        -- e.g. 'BKK', 'PHUKET_CONTRACT'
        destination  TEXT NOT NULL,        -- e.g. 'Bangkok', 'Phuket'
        service_type TEXT NOT NULL,        -- Transfer | Tour | Disposal | Combo | Enroute
        service_name TEXT NOT NULL,
        tour_code    TEXT,
        duration     TEXT,
        includes_vat INTEGER NOT NULL DEFAULT 0,   -- 1 = already includes 7% VAT
        notes        TEXT
    );

    CREATE TABLE IF NOT EXISTS rates (
        id           INTEGER PRIMARY KEY AUTOINCREMENT,
        service_id   INTEGER NOT NULL REFERENCES services(id),
        rate_type    TEXT NOT NULL,        -- Private | SIC
        vehicle      TEXT,                 -- CAR | SUV | VAN | Sedan | ... | null for SIC
        pax_range    TEXT,                 -- e.g. '1-2 pax', '3-9 pax', 'All', null
        pax_category TEXT NOT NULL,        -- Adult | Child | PerVehicle
        price_thb    INTEGER NOT NULL
    );

    CREATE TABLE IF NOT EXISTS zone_surcharges (
        id           INTEGER PRIMARY KEY AUTOINCREMENT,
        service_id   INTEGER NOT NULL REFERENCES services(id),
        zone_name    TEXT NOT NULL,
        surcharge    INTEGER NOT NULL,
        per          TEXT NOT NULL DEFAULT 'trip'   -- 'trip' | 'pax'
    );

    CREATE TABLE IF NOT EXISTS addons (
        id           INTEGER PRIMARY KEY AUTOINCREMENT,
        service_id   INTEGER NOT NULL REFERENCES services(id),
        addon_name   TEXT NOT NULL,
        price_adult  INTEGER,
        price_child  INTEGER
    );
    """)
    conn.commit()


def insert_service(conn, **kw) -> int:
    cur = conn.execute(
        """INSERT INTO services
           (source, destination, service_type, service_name, tour_code,
            duration, includes_vat, notes)
           VALUES (:source,:destination,:service_type,:service_name,
                   :tour_code,:duration,:includes_vat,:notes)""",
        {
            "source":       kw.get("source", ""),
            "destination":  kw.get("destination", ""),
            "service_type": kw.get("service_type", ""),
            "service_name": kw.get("service_name", ""),
            "tour_code":    kw.get("tour_code"),
            "duration":     kw.get("duration"),
            "includes_vat": kw.get("includes_vat", 0),
            "notes":        kw.get("notes"),
        },
    )
    return cur.lastrowid


def insert_rate(conn, service_id, rate_type, vehicle, pax_range, pax_category, price_thb):
    if price_thb is None:
        return
    conn.execute(
        """INSERT INTO rates
           (service_id, rate_type, vehicle, pax_range, pax_category, price_thb)
           VALUES (?,?,?,?,?,?)""",
        (service_id, rate_type, vehicle, pax_range, pax_category, price_thb),
    )


def insert_zone(conn, service_id, zone_name, surcharge, per="trip"):
    if surcharge is None:
        return
    conn.execute(
        "INSERT INTO zone_surcharges (service_id, zone_name, surcharge, per) VALUES (?,?,?,?)",
        (service_id, zone_name, surcharge, per),
    )


def insert_addon(conn, service_id, addon_name, price_adult=None, price_child=None):
    conn.execute(
        "INSERT INTO addons (service_id, addon_name, price_adult, price_child) VALUES (?,?,?,?)",
        (service_id, addon_name, price_adult, price_child),
    )


# ── PARSER 1: BKK/PTY Excel ──────────────────────────────────────────────────
VEHICLE_MAP_BKK = {0: "CAR", 1: "SUV", 2: "VAN"}

def parse_bkk_sheet(conn, xl, sheet_name: str, service_type: str):
    df = pd.read_excel(xl, sheet_name=sheet_name, header=None)
    # row 1 is the header row (SL.No, Service Name, City, Service Type, CAR, SUV, VAN)
    # data starts at row 2
    for _, row in df.iloc[2:].iterrows():
        sl   = str(row.iloc[0]).strip()
        name = clean(str(row.iloc[1]))
        city = clean(str(row.iloc[2]))
        if not sl.isdigit() or not name or name in ("nan", "Service Name"):
            continue
        car = to_int(row.iloc[4])
        suv = to_int(row.iloc[5])
        van = to_int(row.iloc[6])
        if car is None and suv is None and van is None:
            continue

        sid = insert_service(
            conn,
            source="BKK_EXCEL",
            destination=city,
            service_type=service_type,
            service_name=name,
            includes_vat=0,
        )
        for vehicle, price in [("CAR", car), ("SUV", suv), ("VAN", van)]:
            insert_rate(conn, sid, "Private", vehicle, None, "PerVehicle", price)


def parse_bkk_excel(conn):
    print("  Parsing BKK/PTY Excel …")
    xl = pd.ExcelFile(BKK_XL)
    for sheet, stype in [
        ("TRANSFER",    "Transfer"),
        ("TOUR",        "Tour"),
        ("ENROUTE COMBI", "Enroute"),
        ("DISPOSAL",    "Disposal"),
    ]:
        parse_bkk_sheet(conn, xl, sheet, stype)
    conn.commit()
    print("  ✓ BKK Excel done")


# ── PARSER 2: Good Day Vacation Excel ────────────────────────────────────────
def parse_gdv_airport_transfers(conn):
    """Sheet: Transportation Arrival  — airport to hotel zones."""
    df = pd.read_excel(GDV_XL, sheet_name="Transportation Arrival", header=None)
    # row 2 = vehicle type names, row 3 = pax ranges, data rows 5+
    vehicles = ["Sedan/SUV", "D4D Van", "Camry", "Hyundai", "VIP Van", "Alphard"]
    col_offset = 1  # column index where first vehicle price is

    for _, row in df.iloc[5:].iterrows():
        dest = clean(str(row.iloc[0]))
        if not dest or dest in ("nan", "Hotel Zone", "Pier Zone") or "SERVICE" in dest:
            continue

        sid = insert_service(
            conn,
            source="GDV_EXCEL",
            destination="Phuket",
            service_type="Transfer",
            service_name=f"Phuket Airport → {dest}",
            notes="Airport arrival transfer",
            includes_vat=0,
        )
        pax_ranges = ["1-3 pax", "4-8 pax", "1-2 pax", "1-4 pax", "1-5 pax", "1-4 pax"]
        for i, (v, pr) in enumerate(zip(vehicles, pax_ranges)):
            price = to_int(row.iloc[col_offset + i])
            insert_rate(conn, sid, "Private", v, pr, "PerVehicle", price)


def parse_gdv_pickup_dropoff(conn):
    """Sheet: Pick up and Drop off — roundtrip tour transfer per hotel zone."""
    df = pd.read_excel(GDV_XL, sheet_name="Pick up and Drop off ", header=None)
    # row 2 has hotel zone cols: Patong, Karon, Kata, Kamala Laguna, Maikhao Rawai Panwa
    zones = ["Patong Beach", "Karon Beach", "Kata Beach", "Kamala/Laguna", "Maikhao/Rawai/Panwa"]

    for _, row in df.iloc[3:].iterrows():
        dest = clean(str(row.iloc[0]))
        if not dest or dest in ("nan",) or "ROUNDTRIP" in dest or "Destination" in dest:
            continue
        if "PRIVATE TRANSFER" in dest:
            break  # reached the next section

        sid = insert_service(
            conn,
            source="GDV_EXCEL",
            destination="Phuket",
            service_type="Transfer",
            service_name=f"Phuket R/T Hotel Pickup: {dest}",
            notes="Roundtrip hotel pickup/dropoff",
            includes_vat=0,
        )
        col = 1
        for zone in zones:
            car_price = to_int(row.iloc[col])
            van_price = to_int(row.iloc[col + 1])
            if car_price:
                insert_rate(conn, sid, "Private", "CAR", None, "PerVehicle", car_price)
                insert_zone(conn, sid, zone, 0)  # zone info embedded in base price here
            if van_price:
                insert_rate(conn, sid, "Private", "VAN", None, "PerVehicle", van_price)
            col += 2


def parse_gdv_intercity(conn):
    """Sheet: Pick up and Drop off — intercity transfers at the bottom of the sheet."""
    df = pd.read_excel(GDV_XL, sheet_name="Pick up and Drop off ", header=None)
    in_intercity = False
    for _, row in df.iterrows():
        first = clean(str(row.iloc[0]))
        if "PRIVATE TRANSFER FROM PHUKET" in first:
            in_intercity = True
            continue
        if not in_intercity or not first or first == "nan":
            continue
        name = first
        car = to_int(row.iloc[8]) if len(row) > 8 else None
        suv = to_int(row.iloc[9]) if len(row) > 9 else None
        van = to_int(row.iloc[10]) if len(row) > 10 else None
        if car is None and suv is None and van is None:
            continue

        # Determine destination
        dest = "Phuket"
        if "Krabi" in name:
            dest = "Krabi"
        elif "Khao Lak" in name or "Natai" in name or "Phang Nga" in name:
            dest = "Phang Nga"

        sid = insert_service(
            conn,
            source="GDV_EXCEL",
            destination=dest,
            service_type="Transfer",
            service_name=name,
            includes_vat=0,
        )
        for vehicle, price in [("CAR", car), ("SUV", suv), ("VAN", van)]:
            insert_rate(conn, sid, "Private", vehicle, None, "PerVehicle", price)


def parse_gdv_samui(conn):
    """Sheet: Samui Transfers & TOUR EXCURSI."""
    df = pd.read_excel(GDV_XL, sheet_name="Samui Transfers & TOUR  EXCURSI", header=None)
    for _, row in df.iterrows():
        from_loc = clean(str(row.iloc[0]))
        to_loc   = clean(str(row.iloc[1]))
        price    = to_int(row.iloc[3]) if len(row) > 3 else None

        if not from_loc or from_loc in ("nan", "From") or "Transfer Rate" in from_loc:
            continue
        if price is None:
            continue

        sid = insert_service(
            conn,
            source="GDV_EXCEL",
            destination="Samui",
            service_type="Transfer",
            service_name=f"{from_loc} → {to_loc}",
            includes_vat=0,
        )
        insert_rate(conn, sid, "Private", "VAN", None, "PerVehicle", price)


def parse_gdv_krabi_excursion(conn):
    """Sheet: Tour Excursion in Krabi."""
    df = pd.read_excel(GDV_XL, sheet_name="Tour  Excursion in Krabi", header=None)
    for _, row in df.iloc[3:].iterrows():
        name = clean(str(row.iloc[0]))
        if not name or name == "nan" or "TOUR" in name.upper()[:4]:
            continue
        adult_ex  = to_int(row.iloc[2]) if len(row) > 2 else None
        child_ex  = to_int(row.iloc[3]) if len(row) > 3 else None
        adult_inc = to_int(row.iloc[4]) if len(row) > 4 else None
        child_inc = to_int(row.iloc[5]) if len(row) > 5 else None

        if adult_ex is None and adult_inc is None:
            continue

        sid = insert_service(
            conn,
            source="GDV_EXCEL",
            destination="Krabi",
            service_type="Tour",
            service_name=name,
            notes="SIC join trip",
            includes_vat=0,
        )
        if adult_ex:
            insert_rate(conn, sid, "SIC", None, None, "Adult", adult_ex)
        if child_ex:
            insert_rate(conn, sid, "SIC", None, None, "Child", child_ex)
        if adult_inc:
            insert_rate(conn, sid, "SIC", None, None, "Adult", adult_inc)
        if child_inc:
            insert_rate(conn, sid, "SIC", None, None, "Child", child_inc)


def parse_gdv_combo(conn):
    """Sheet: Combo Rates."""
    df = pd.read_excel(GDV_XL, sheet_name="Combo Rates", header=None)
    for _, row in df.iloc[4:].iterrows():
        name = clean(str(row.iloc[0]))
        if not name or name == "nan" or "COMBO" in name.upper()[:5]:
            continue
        car = to_int(row.iloc[2])
        van = to_int(row.iloc[3])
        if car is None and van is None:
            continue
        sid = insert_service(
            conn,
            source="GDV_EXCEL",
            destination="Phuket",
            service_type="Combo",
            service_name=name,
            includes_vat=0,
        )
        insert_rate(conn, sid, "Private", "CAR", None, "PerVehicle", car)
        insert_rate(conn, sid, "Private", "VAN", None, "PerVehicle", van)


def parse_gdv_excel(conn):
    print("  Parsing Good Day Vacation Excel …")
    parse_gdv_airport_transfers(conn)
    parse_gdv_pickup_dropoff(conn)
    parse_gdv_intercity(conn)
    parse_gdv_samui(conn)
    parse_gdv_krabi_excursion(conn)
    parse_gdv_combo(conn)
    conn.commit()
    print("  ✓ GDV Excel done")


# ── PARSER 3: Word docs (Phuket / Krabi / Samui) ─────────────────────────────

def get_unique_row(row) -> list[str]:
    """Deduplicate a merged Word table row into unique non-empty strings."""
    seen, result = [], []
    for cell in row.cells:
        t = clean(cell.text)
        if t and t not in seen:
            seen.append(t)
            result.append(t)
    return result


def extract_tour_code(text: str) -> str | None:
    m = re.search(r"SML\s*\d+", text, re.IGNORECASE)
    return m.group(0).replace(" ", " ").strip() if m else None


def parse_contract_docx(conn, docx_path: Path, destination: str, source_tag: str, includes_vat: int):
    """
    Parse a contract .docx file.
    Each tour is a multi-row table block.
    Row 0 = tour code title (merged)
    Row 1 = description (merged)
    Row 2 = column headers section labels
    Row 3 = vehicle / pax_category sub-headers
    Row 4..N = data rows with zone name + prices
    """
    print(f"  Parsing {docx_path.name} …")
    doc = Document(docx_path)

    for tbl in doc.tables:
        rows = tbl.rows
        if len(rows) < 4:
            continue

        # ── Find tour code row ──────────────────────────────────────────────
        row0_text = clean(rows[0].cells[0].text)
        tour_code = extract_tour_code(row0_text)
        if not tour_code:
            # Check if it's a transfer table
            if "TRANSFER" not in row0_text.upper():
                continue
            tour_name = row0_text
        else:
            tour_name = row0_text

        # ── Determine columns from header rows ──────────────────────────────
        # We scan rows 2 and 3 to find vehicle columns and SIC columns
        hdr2 = get_unique_row(rows[2]) if len(rows) > 2 else []
        hdr3 = get_unique_row(rows[3]) if len(rows) > 3 else []

        # Map raw cell indices to (rate_type, vehicle, pax_category)
        # Strategy: iterate actual cell list (with duplicates = merged) to get column indices
        col_map = {}  # col_index → (rate_type, vehicle, pax_category, pax_range)

        raw2 = [clean(c.text) for c in rows[2].cells] if len(rows) > 2 else []
        raw3 = [clean(c.text) for c in rows[3].cells] if len(rows) > 3 else []

        current_section = "Private"
        current_vehicle = None

        for ci, (h2, h3) in enumerate(zip(raw2, raw3)):
            if "SIC" in h2.upper():
                current_section = "SIC"
            if "PRIVATE" in h2.upper():
                current_section = "Private"

            # vehicle detection in h3
            for v in ["CAR", "SUV", "VAN"]:
                if v in h3.upper() and not any(x in h3.upper() for x in ["GUIDE", "PARK", "LUNCH", "ADULT", "CHILD"]):
                    current_vehicle = v
                    break

            if "ADULT" in h3.upper() and "GUIDE" not in h3.upper():
                pax_range = None
                m = re.search(r"(\d+-\d+)\s*PAX", h3, re.IGNORECASE)
                if m:
                    pax_range = m.group(1) + " pax"
                col_map[ci] = (current_section, None, "Adult", pax_range)

            elif "CHILD" in h3.upper() and "GUIDE" not in h3.upper():
                pax_range = None
                m = re.search(r"(\d+-\d+)\s*PAX", h3, re.IGNORECASE)
                if m:
                    pax_range = m.group(1) + " pax"
                col_map[ci] = (current_section, None, "Child", pax_range)

            elif current_section == "Private" and current_vehicle:
                pax_range = None
                m = re.search(r"(\d+-\d+)\s*PAX", h3, re.IGNORECASE)
                if m:
                    pax_range = m.group(1) + " pax"
                col_map[ci] = ("Private", current_vehicle, "PerVehicle", pax_range)

        if not col_map:
            continue

        # ── Parse data rows ─────────────────────────────────────────────────
        # Data rows start at row 4
        data_rows = rows[4:]
        prev_zone = "All areas"
        services_in_table: dict[str, int] = {}  # zone → service_id

        for dr in data_rows:
            raw = [clean(c.text) for c in dr.cells]
            if not raw:
                continue

            zone_candidate = raw[0]

            # Skip footnote / condition rows
            skip_keywords = ["SIC RATES", "PRIVATE RATES", "EXTRA CHARGE", "PICK UP",
                             "CHILD RATES", "NOTE", "ENGLISH", "HOTELS IN", "AVAILABLE"]
            if any(kw in zone_candidate.upper() for kw in skip_keywords):
                continue

            # Check if any price columns have numeric values
            has_prices = any(to_int(raw[ci]) is not None for ci in col_map if ci < len(raw))
            if not has_prices:
                continue

            zone = zone_candidate if zone_candidate and zone_candidate != "nan" else prev_zone
            if zone:
                prev_zone = zone

            # Create or reuse service for this zone
            if zone not in services_in_table:
                sid = insert_service(
                    conn,
                    source=source_tag,
                    destination=destination,
                    service_type="Tour",
                    service_name=tour_name,
                    tour_code=tour_code,
                    includes_vat=includes_vat,
                    notes=f"Zone: {zone}",
                )
                services_in_table[zone] = sid

            sid = services_in_table[zone]

            for ci, (rtype, vehicle, pax_cat, pax_range) in col_map.items():
                if ci >= len(raw):
                    continue
                price = to_int(raw[ci])
                insert_rate(conn, sid, rtype, vehicle, pax_range, pax_cat, price)

    # ── Parse transfer tables at end of docx (Krabi / Samui have these) ─────
    _parse_transfer_tables_docx(conn, doc, destination, source_tag, includes_vat)

    conn.commit()
    print(f"  ✓ {docx_path.name} done")


def _parse_transfer_tables_docx(conn, doc: Document, destination: str, source_tag: str, includes_vat: int):
    """Pick up any TRANSFER RATES tables in the docx (appear in Krabi & Samui)."""
    for tbl in doc.tables:
        row0 = clean(tbl.rows[0].cells[0].text) if tbl.rows else ""
        if "TRANSFER RATES" not in row0.upper():
            continue

        # Find price columns — look for PAX range headers
        raw2 = [clean(c.text) for c in tbl.rows[2].cells] if len(tbl.rows) > 2 else []
        raw3 = [clean(c.text) for c in tbl.rows[3].cells] if len(tbl.rows) > 3 else []
        raw1 = [clean(c.text) for c in tbl.rows[1].cells] if len(tbl.rows) > 1 else []

        # simple 2-col case: FROM, TO, PRICE, PRICE
        for dr in tbl.rows[3:]:
            raw = [clean(c.text) for c in dr.cells]
            uq  = list(dict.fromkeys(r for r in raw if r and r != "nan"))
            if len(uq) < 3:
                continue
            from_loc = uq[0]
            to_loc   = uq[1] if len(uq) > 1 else ""
            skip_kw  = ["FROM", "DESTINATION", "RATE", "SIC", "TRANSFER", "REMARK"]
            if any(k in from_loc.upper() for k in skip_kw):
                continue
            # Try to find a price
            price = None
            for v in uq[2:]:
                price = to_int(v)
                if price and price > 100:
                    break
            if price is None:
                continue

            sid = insert_service(
                conn,
                source=source_tag,
                destination=destination,
                service_type="Transfer",
                service_name=f"{from_loc} → {to_loc}",
                includes_vat=includes_vat,
            )
            insert_rate(conn, sid, "Private", "VAN", None, "PerVehicle", price)


def parse_all_docx(conn):
    parse_contract_docx(conn, PHK_DOCX, "Phuket", "PHUKET_CONTRACT", includes_vat=0)
    parse_contract_docx(conn, KRB_DOCX, "Krabi",  "KRABI_CONTRACT",  includes_vat=0)
    parse_contract_docx(conn, SMU_DOCX, "Samui",  "SAMUI_CONTRACT",  includes_vat=1)


# ── main ─────────────────────────────────────────────────────────────────────
def main():
    if DB_PATH.exists():
        DB_PATH.unlink()
        print(f"Removed existing {DB_PATH}")

    conn = sqlite3.connect(DB_PATH)
    conn.execute("PRAGMA foreign_keys = ON")
    init_db(conn)

    print("\nParsing source files …")
    parse_bkk_excel(conn)
    parse_gdv_excel(conn)
    parse_all_docx(conn)

    # ── Summary ──────────────────────────────────────────────────────────────
    print("\n── Database summary ─────────────────────────────────────────────")
    for tbl in ["services", "rates", "zone_surcharges", "addons"]:
        n = conn.execute(f"SELECT COUNT(*) FROM {tbl}").fetchone()[0]
        print(f"  {tbl:<20} {n:>5} rows")

    print("\n  Services by destination:")
    for row in conn.execute(
        "SELECT destination, COUNT(*) FROM services GROUP BY destination ORDER BY destination"
    ):
        print(f"    {row[0]:<20} {row[1]:>4}")

    print("\n  Services by type:")
    for row in conn.execute(
        "SELECT service_type, COUNT(*) FROM services GROUP BY service_type ORDER BY service_type"
    ):
        print(f"    {row[0]:<20} {row[1]:>4}")

    conn.close()
    print(f"\n✓ Done — {DB_PATH}")


if __name__ == "__main__":
    main()
